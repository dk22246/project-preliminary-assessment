"""统一的《招商项目整体落地研判报告》Word 样式与表格构建组件。"""
from __future__ import annotations

from pathlib import Path
import subprocess
import tempfile
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PAGE = {"top": 2.5, "bottom": 2.5, "left": 2.8, "right": 2.5, "header": 1.5, "footer": 1.5}
TABLE_HEADER_FILL = "E7E6E6"
TABLE_BORDER = "595959"


def _set_font(style_or_run, *, chinese: str, size: float, bold: bool = False) -> None:
    style_or_run.font.name = "Arial"
    style_or_run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style_or_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style_or_run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    style_or_run.font.size = Pt(size)
    style_or_run.font.bold = bold


def configure_report_document(doc, report_short_name: str) -> None:
    """应用统一A4版式、原生标题样式及页眉页脚。"""
    for section in doc.sections:
        configure_section(section)
        section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    _set_font(normal, chinese="宋体", size=12)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    _set_heading_style(doc.styles["Heading 1"], 16, 12, 6, chinese="黑体")
    _set_heading_style(doc.styles["Heading 2"], 14, 8, 4, chinese="黑体")
    _set_heading_style(doc.styles["Heading 3"], 12, 6, 3, chinese="宋体")

    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(report_short_name)
    _set_font(run, chinese="宋体", size=10.5)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)
    _set_font(footer.runs[0], chinese="宋体", size=10.5)
    # Cover intentionally has no page furniture.
    doc.sections[0].first_page_header.paragraphs[0].text = ""
    doc.sections[0].first_page_footer.paragraphs[0].text = ""


def configure_section(section, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width, section.page_height = Cm(29.7), Cm(21)
    else:
        section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = Cm(PAGE["top"])
    section.bottom_margin = Cm(PAGE["bottom"])
    section.left_margin = Cm(PAGE["left"])
    section.right_margin = Cm(PAGE["right"])
    section.header_distance = Cm(PAGE["header"])
    section.footer_distance = Cm(PAGE["footer"])


def _set_heading_style(style, size: float, before: float, after: float, *, chinese: str) -> None:
    _set_font(style, chinese=chinese, size=size, bold=True)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.keep_together = True


def add_heading(doc, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text: str, *, centered: bool = False, role: str = "body"):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    if role == "source":
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        _set_font(run, chinese="宋体", size=10.5)
    else:
        paragraph.add_run(text)
    return paragraph


def add_cover_line(doc, text: str, *, title: bool = False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8 if title else 6)
    run = paragraph.add_run(text)
    _set_font(run, chinese="黑体" if title else "宋体", size=22 if title else 12, bold=title)
    return paragraph


def add_native_toc_with_cache(paragraph, entries: Sequence[tuple[str, int]]) -> None:
    """插入可更新的Word原生TOC域，并给无自动更新环境提供可见缓存。"""
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    for index, (label, page) in enumerate(entries):
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.set(qn("xml:space"), "preserve")
        text.text = f"{label}{' ' * max(2, 44 - len(label))}{page}"
        run.append(text)
        field.append(run)
        if index < len(entries) - 1:
            br_run = OxmlElement("w:r")
            br_run.append(OxmlElement("w:br"))
            field.append(br_run)
    paragraph._p.append(field)


def try_update_fields_with_word(path: Path, timeout_seconds: int = 30) -> bool:
    """用本机 Word 刷新 TOC、页码及交叉引用；不可用时保留可见 TOC 缓存。"""
    script = r'''
param([string]$DocumentPath)
$word = $null
$document = $null
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $document = $word.Documents.Open($DocumentPath, $false, $false)
    foreach ($toc in $document.TablesOfContents) { $toc.Update() }
    foreach ($field in $document.Fields) { $field.Update() }
    $document.Save()
    exit 0
}
catch {
    Write-Error $_
    exit 1
}
finally {
    if ($document -ne $null) { $document.Close(0) }
    if ($word -ne $null) { $word.Quit() }
}
'''
    temporary = tempfile.NamedTemporaryFile(mode="w", suffix=".ps1", encoding="utf-8", delete=False)
    try:
        temporary.write(script)
        temporary.close()
        result = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-ExecutionPolicy", "Bypass", "-File", temporary.name, str(path)],
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            check=False,
        )
        if result.returncode == 0:
            _ensure_update_fields_on_open(path)
            return True
        return False
    except (OSError, subprocess.TimeoutExpired):
        return False
    finally:
        Path(temporary.name).unlink(missing_ok=True)


def _ensure_update_fields_on_open(path: Path) -> None:
    """Word 刷新后重新写回 updateFields，保证后续打开时仍会自动更新。"""
    document = Document(path)
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    document.save(path)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("第 ")
    _set_font(run, chinese="宋体", size=10.5)
    for kind, text in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
        node = OxmlElement("w:fldChar" if kind else "w:instrText" if text == " PAGE " else "w:t")
        if kind:
            node.set(qn("w:fldCharType"), kind)
        elif text == " PAGE ":
            node.set(qn("xml:space"), "preserve")
        node.text = text
        field_run = OxmlElement("w:r")
        field_run.append(node)
        paragraph._p.append(field_run)
    run = paragraph.add_run(" 页")
    _set_font(run, chinese="宋体", size=10.5)


def add_standard_table(
    doc,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths_cm: Sequence[float],
    *,
    centered_columns: Iterable[int] = (),
    numeric_columns: Iterable[int] = (),
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_geometry(table, widths_cm)
    _set_repeat_header(table.rows[0])
    centered = set(centered_columns)
    numeric = set(numeric_columns)
    for column, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        _shade(cell, TABLE_HEADER_FILL)
        _write_cell(cell, label, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for values in rows:
        row = table.add_row()
        _prevent_row_split(row)
        for column, (cell, value) in enumerate(zip(row.cells, values)):
            alignment = WD_ALIGN_PARAGRAPH.RIGHT if column in numeric else WD_ALIGN_PARAGRAPH.CENTER if column in centered else WD_ALIGN_PARAGRAPH.LEFT
            _write_cell(cell, str(value), alignment=alignment)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Cm(0)
    spacer.paragraph_format.space_after = Pt(4)
    return table


def _write_cell(cell, text: str, *, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _set_font(run, chinese="宋体", size=10.5, bold=bold)


def _set_table_geometry(table, widths_cm: Sequence[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total = sum(round(width * 567) for width in widths_cm)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    for grid, width in zip(table._tbl.tblGrid.gridCol_lst, widths_cm):
        grid.set(qn("w:w"), str(round(width * 567)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            cell._tc.get_or_add_tcPr().tcW.set(qn("w:w"), str(round(width * 567)))
            cell._tc.tcPr.tcW.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
