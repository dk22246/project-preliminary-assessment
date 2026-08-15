"""Optional editable Word renderer from the same report-data.json used by HTML/PDF."""
from __future__ import annotations

import argparse
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm

from report_core import financial_change_notes, financial_headers, load_data, validate_report_data, validate_text
from word_report_builder import add_body, add_cover_line, add_heading, add_native_toc_with_cache, add_standard_table, configure_report_document, try_update_fields_with_word


def rows(items, fields):
    return [[str(item.get(field, "未公开披露")) for field in fields] for item in items]


def h1(doc, title):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, title, 1)


def add_equity_conflict_disclosures(doc, items):
    if not items:
        return
    add_heading(doc, "股权数据差异说明", 3)
    for item in items:
        status = "已核实" if item.get("status") == "resolved" else "待核实"
        add_body(doc, f"{item.get('title', '未命名差异')}（{status}）")
        add_body(doc, f"差异情况：{item.get('difference', '未说明')}")
        add_body(doc, f"可能原因：{item.get('reason', '未说明')}")
        add_body(doc, f"本报告处理口径：{item.get('adopted_basis', '未说明')}")
        add_body(doc, f"对招商判断的影响：{item.get('impact', '未说明')}")
        add_body(doc, f"后续核实：{item.get('next_action', '未说明')}")
        add_body(doc, "证据来源：" + "、".join(str(source_id) for source_id in item.get("evidence_source_ids", [])))


def add_equity_evidence_summary(doc, equity):
    summary = equity.get("evidence_summary", {})
    add_body(doc, f"股权来源：{summary.get('display_source_title', '需补充股权来源')}（{summary.get('display_source_id', 'E?')}），数据时点{summary.get('as_of_date', '需补充')}。")


def encouraged_industry_rows(assessment):
    labels = {"direct_match": "明确符合", "potential_match": "存在相近可能", "no_match": "暂未发现明确匹配"}
    rendered = []
    for item in assessment.get("business_assessments", []):
        matched = "；".join(
            f"{candidate.get('catalog_item_no', '—')}．{candidate.get('catalog_item', '未注明条目')}｜{candidate.get('detailed_item', '未注明细化目录')}"
            for candidate in item.get("matched_items", [])
        ) or "完成三条目录路径检索后，未发现可合理对应的具体条目"
        rendered.append([item.get("activity_name", item.get("business", "未命名经营活动")), labels.get(item.get("judgment"), "研究未完成"), matched, item.get("reason", "未说明"), item.get("verification_needed", "无")])
    return rendered


def add_industry_chain(doc, chain):
    add_body(doc, "产业链定位：" + chain.get("positioning", "需企业补充"))
    labels = {"upstream": "上游", "midstream": "中游", "downstream": "下游"}
    rendered = []
    for stage in chain.get("stages", []):
        enterprises = "、".join(item.get("name", "") for item in stage.get("representative_enterprises", [])) or "本轮未发现可靠代表企业"
        rendered.append([labels.get(stage.get("stage"), "环节"), stage.get("title", "未命名"), "；".join(stage.get("activities", [])), enterprises])
    add_standard_table(doc, ["环节", "定位", "核心活动", "代表企业"], rendered, [1.6, 3.2, 6.0, 5.1])
    add_body(doc, "区域产业生态：" + chain.get("regional_ecosystem", {}).get("summary", "本轮未发现可靠资料"))


def policy_match_rows(policies):
    """Keep the editable Word table aligned with the HTML decision table."""
    groups = {}
    for item in policies:
        group = str(item.get("report_group") or item.get("source_id") or item.get("name"))
        groups.setdefault(group, []).append(item)
    rendered = []
    for group in groups.values():
        item = group[0]
        references = "；".join(f"{policy.get('document_number', '未公开披露')}（{policy.get('source_id', 'P')}）" for policy in group)
        rendered.append([
            f"{item.get('report_title') or item.get('name', '未命名政策')}\n政策依据：{references}",
            str(item.get("report_reason", "未公开披露")),
        ])
    return rendered


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("report_data")
    parser.add_argument("--out", required=True)
    parser.add_argument("--equity-image", required=True)
    args = parser.parse_args()
    data = load_data(args.report_data)
    errors = validate_report_data(data) + validate_text(data)
    if errors:
        raise SystemExit("\n".join(errors))
    image = Path(args.equity_image)
    if not image.exists():
        raise SystemExit("缺少股权图 PNG，Word 不得用文本框替代图形")
    doc = Document()
    configure_report_document(doc, data["meta"].get("report_short_name", "招商项目整体落地研判报告"))
    add_cover_line(doc, data["meta"]["report_title"], title=True)
    add_cover_line(doc, f"编制单位：{data['meta'].get('unit', '三亚中央商务区招商研判组')}")
    add_cover_line(doc, f"编制日期：{data['meta']['generated_at']}")
    doc.add_page_break(); add_heading(doc, "目录", 1)
    add_native_toc_with_cache(doc.add_paragraph(), [(x, 1) for x in ("一、企业基本情况", "二、近三年经营数据", "三、风险与合规情况", "四、三亚落地业务及落地方式", "五、企业政策匹配", "六、综合评估", "七、参考资料")])
    h1(doc, "一、企业基本情况")
    add_heading(doc, "（一）企业主体认定与企业概况", 2)
    entity = data["entity_resolution"]
    overview = data["enterprise_overview"]
    entity_rows = [["企业主体", entity.get("legal_entity", "未公开披露")], ["成立时间", overview["established_at"]], ["注册地", overview["registered_location"]], ["企业性质", overview["listing_status"]], ["主营业务", overview["main_business"]], ["分析口径", entity.get("financial_scope", "未公开披露")]]
    add_standard_table(doc, ["基本事项", "企业情况"], entity_rows, [4.2, 11.7])
    add_body(doc, "企业概况：" + overview["profile"])
    add_body(doc, "经营表现：" + overview["operating_summary"])
    add_body(doc, "员工规模：" + overview["employee_scale"])
    industry = data.get("industry_position", {})
    if isinstance(industry, dict):
        industry_text = f"{industry.get('category', '相关行业')}品类定位：{industry.get('position', '未取得可靠排名')}；统计时点：{industry.get('period', '未注明时点')}。{industry.get('statement', '需企业补充')}"
    else:
        industry_text = str(industry or "需企业补充")
    add_body(doc, "行业地位：" + industry_text)
    add_heading(doc, "（二）股权架构拆解", 2); doc.add_picture(str(image), width=Cm(15)); add_body(doc, entity.get("equity_summary", "需企业补充")); add_equity_evidence_summary(doc, data["equity"]); add_equity_conflict_disclosures(doc, data["equity"].get("conflict_disclosures", []))
    add_heading(doc, "（三）主要业务及产品拆解", 2); add_standard_table(doc, ["业务板块", "主要产品或服务", "主要承载主体", "销售渠道", "国内外业务布局"], rows(data["businesses"], ("segment", "products", "entity", "sales_channels", "footprint")), [2.4, 3.5, 3.0, 3.8, 3.2])
    add_heading(doc, "（四）海南自由贸易港鼓励类产业目录匹配", 2); add_body(doc, "总体判断：" + data["encouraged_industry_assessment"].get("summary", "未说明")); add_standard_table(doc, ["企业具体经营活动", "匹配结论", "对应目录条目", "判断依据", "相近可能或待核事项"], encouraged_industry_rows(data["encouraged_industry_assessment"]), [2.5, 1.7, 4.0, 4.8, 3.0])
    add_heading(doc, "（五）产业链上下游", 2); add_industry_chain(doc, data["industry_chain"])
    h1(doc, "二、近三年经营数据")
    headers = financial_headers(data["meta"])
    add_heading(doc, "（一）营业收入、利润、纳税及政府补助情况", 2); add_standard_table(doc, headers, rows(data["financials"], ("year", "revenue", "revenue_change", "profit", "profit_change", "tax_value", "tax_basis", "government_support", "source")), [1.1, 1.55, 0.9, 1.45, 0.9, 1.5, 1.5, 1.6, 1.0])
    change_notes = financial_change_notes(data["financials"])
    if change_notes:
        add_body(doc, "注：" + "；".join(change_notes))
    add_heading(doc, "政府补助及财政支持明细表", 3); add_standard_table(doc, ["年度", "名称", "发放部门", "金额", "用途", "附带条件", "来源"], rows(data.get("government_support", []), ("year", "name", "department", "amount", "purpose", "conditions", "source")) or [["—", "本轮公开检索未发现可确认明细", "—", "—", "需企业补充", "需企业补充", "—"]], [1.1, 2.5, 2.0, 1.3, 3.0, 3.0, 1.0])
    add_heading(doc, "（二）经营数据分析", 2); add_body(doc, data.get("financial_analysis", "需企业补充"))
    h1(doc, "三、风险与合规情况")
    add_body(doc, data["risks"].get("summary", "截至公开检索未发现重大记录，仍需企业及主管部门核验。"))
    h1(doc, "四、三亚落地业务及落地方式"); add_standard_table(doc, ["建议落地业务", "企业现有事实基础", "三亚具体承接方式", "可形成的业务及价值", "可行性"], rows(data["landing_businesses"], ("business", "fact_basis", "sanya_path", "value", "feasibility")), [3.0, 4.0, 4.0, 4.0, 1.0])
    h1(doc, "五、企业政策匹配")
    add_heading(doc, "（一）重点政策匹配清单", 2)
    add_body(doc, "本节依据企业已公开的业务、组织和境内外布局，展示在三亚承接相邻经营活动时可重点沟通的现行政策或办理工具。完整检索、失效政策及排除理由保留在后台台账。")
    add_standard_table(doc, ["匹配政策或工具", "匹配原因"], policy_match_rows(data["policies"]), [5.2, 10.7])
    h1(doc, "六、综合评估"); add_body(doc, data.get("comprehensive_assessment", "需企业补充"))
    h1(doc, "七、参考资料"); add_standard_table(doc, ["编号", "类型", "资料名称", "发布主体", "日期", "网址或文件定位", "使用位置"], rows(data["sources"], ("id", "type", "name", "issuer", "date", "location", "used_in")), [1.0, 1.1, 3.0, 2.3, 1.4, 5.0, 2.0])
    out = Path(args.out); out.parent.mkdir(parents=True, exist_ok=True); doc.save(out)
    try_update_fields_with_word(out); print(out); return 0


if __name__ == "__main__":
    raise SystemExit(main())
